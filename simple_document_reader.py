"""
ПРОСТОЙ ЧИТАТЕЛЬ ДОКУМЕНТОВ
Читает PDF, Word, Excel, текстовые файлы
"""

import os
import PyPDF2
from docx import Document
import openpyxl

class SimpleDocumentReader:
    """Простой читатель документов"""
    
    def read_file(self, file_path):
        """Читает любой документ"""
        if not os.path.exists(file_path):
            return f"❌ Файл не найден: {file_path}"
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_ext == '.pdf':
                return self.read_pdf(file_path)
            elif file_ext == '.docx':
                return self.read_docx(file_path)
            elif file_ext in ['.xlsx', '.xls']:
                return self.read_excel(file_path)
            elif file_ext == '.txt':
                return self.read_text(file_path)
            else:
                return f"❌ Неподдерживаемый формат: {file_ext}"
                
        except Exception as e:
            return f"❌ Ошибка чтения файла: {e}"
    
    def read_pdf(self, file_path):
        """Читает PDF файл"""
        text = ""
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            # Информация о документе
            num_pages = len(pdf_reader.pages)
            text += f"📄 PDF документ: {os.path.basename(file_path)}\n"
            text += f"📖 Страниц: {num_pages}\n\n"
            
            # Чтение текста со страниц
            for page_num in range(min(num_pages, 5)):  # Только первые 5 страниц
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                
                if page_text.strip():
                    text += f"--- Страница {page_num + 1} ---\n"
                    text += page_text[:500] + "...\n\n" if len(page_text) > 500 else page_text + "\n\n"
        
        return text if text else "❌ Не удалось извлечь текст из PDF"
    
    def read_docx(self, file_path):
        """Читает Word документ"""
        text = ""
        
        doc = Document(file_path)
        
        # Информация о документе
        text += f"📄 Word документ: {os.path.basename(file_path)}\n"
        text += f"📝 Параграфов: {len(doc.paragraphs)}\n\n"
        
        # Чтение параграфов
        for i, para in enumerate(doc.paragraphs[:20]):  # Только первые 20 параграфов
            if para.text.strip():
                text += f"{para.text}\n"
        
        return text if text else "❌ Документ пуст"
    
    def read_excel(self, file_path):
        """Читает Excel файл"""
        text = ""
        
        wb = openpyxl.load_workbook(file_path, data_only=True)
        
        text += f"📊 Excel файл: {os.path.basename(file_path)}\n"
        text += f"📑 Листов: {len(wb.sheetnames)}\n\n"
        
        # Чтение каждого листа
        for sheet_name in wb.sheetnames[:3]:  # Только первые 3 листа
            ws = wb[sheet_name]
            
            text += f"--- Лист: {sheet_name} ---\n"
            
            # Чтение первых 10 строк и 5 столбцов
            for row in ws.iter_rows(min_row=1, max_row=10, max_col=5, values_only=True):
                row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                text += f"{row_text}\n"
            
            text += "\n"
        
        return text
    
    def read_text(self, file_path):
        """Читает текстовый файл"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            return f"📝 Текстовый файл: {os.path.basename(file_path)}\n\n{content[:1000]}..."
        except:
            try:
                with open(file_path, 'r', encoding='cp1251') as file:
                    content = file.read()
                return f"📝 Текстовый файл: {os.path.basename(file_path)}\n\n{content[:1000]}..."
            except Exception as e:
                return f"❌ Ошибка чтения текстового файла: {e}"
    
    def batch_read(self, folder_path):
        """Читает все документы в папке"""
        if not os.path.exists(folder_path):
            return "❌ Папка не найдена"
        
        results = []
        
        for filename in os.listdir(folder_path):
            file_path = os.path.join(folder_path, filename)
            
            if os.path.isfile(file_path):
                print(f"📖 Читаю: {filename}")
                content = self.read_file(file_path)
                
                results.append({
                    'filename': filename,
                    'content': content[:500] + "..." if len(content) > 500 else content
                })
        
        return results

# Пример использования
if __name__ == "__main__":
    print("📖 ТЕСТ ЧИТАТЕЛЯ ДОКУМЕНТОВ")
    print("=" * 40)
    
    reader = SimpleDocumentReader()
    
    # Создаем тестовые файлы
    print("\n📝 Создаю тестовые файлы...")
    
    # Тестовый текстовый файл
    with open("test_document.txt", "w", encoding="utf-8") as f:
        f.write("Это тестовый текстовый файл.\n")
        f.write("Елена может читать такие файлы.\n")
        f.write("Вот и всё!")
    
    print("✅ Созданы тестовые файлы")
    print("\n📖 Читаю файлы...")
    
    # Чтение файлов
    if os.path.exists("test_document.txt"):
        content = reader.read_file("test_document.txt")
        print("\n" + content)
    
    print("\n📚 Чем я могу читать:")
    print("• PDF файлы (нужен PyPDF2)")
    print("• Word документы (нужен python-docx)")
    print("• Excel файлы (нужен openpyxl)")
    print("• Текстовые файлы")
    
    print("\n💡 Пример использования:")
    print("reader = SimpleDocumentReader()")
    print('content = reader.read_file("ваш_файл.pdf")')
    print('print(content)')
    
    # Уборка
    if os.path.exists("test_document.txt"):
        os.remove("test_document.txt")
        print("\n🧹 Удалены тестовые файлы")